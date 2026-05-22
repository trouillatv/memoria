# Convertit un .md de docs/ en .docx.
#   Usage : python scripts/dev/md-to-docx.py [BASENAME]
#   - sans argument → MODE_EMPLOI (+ copie public/manuel.docx)
#   - avec argument → docs/{BASENAME}.md → docs/{BASENAME}.docx
#     ex : python scripts/dev/md-to-docx.py COMPRENDRE_ARCHITECTURE
# Parser markdown minimaliste : headers # à ####, listes -, tables |...|,
# bold **, italic *, inline code `...`, lignes ---.
# python-docx 1.2.0 — pas de dépendance externe à markdown.

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]

# Basename cible : argument CLI ou MODE_EMPLOI par défaut.
BASENAME = sys.argv[1] if len(sys.argv) > 1 else "MODE_EMPLOI"
SRC = ROOT / "docs" / f"{BASENAME}.md"
DST = ROOT / "docs" / f"{BASENAME}.docx"

if not SRC.exists():
    print(f"[ERREUR] Source introuvable : {SRC}")
    sys.exit(1)

# Fallback si DST verrouillé (Word ouvert) — on écrit à côté.
if DST.exists():
    try:
        import os
        fd = os.open(str(DST), os.O_RDWR)
        os.close(fd)
    except (PermissionError, OSError):
        import datetime
        ts = datetime.datetime.now().strftime("%Y%m%d-%H%M")
        DST = ROOT / "docs" / f"{BASENAME}.regen-{ts}.docx"
        print(f"[INFO] {BASENAME}.docx verrouille - ecriture dans {DST.name}")

BRAND = RGBColor(0xC0, 0x39, 0x2B)   # rouge bordeaux MemorIA
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
CODE_BG = RGBColor(0xF4, 0xF4, 0xF4)


def add_run_with_inline(paragraph, text):
    """Pose du texte avec **bold**, *italic*, `code` reconnus."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        elif token.startswith('*'):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_table(doc, lines):
    """Crée un tableau Word depuis des lignes markdown |a|b|c|."""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    # Ligne 2 du markdown = séparateur ---|---|--- → on l'ignore
    if len(rows) >= 2 and all(re.match(r'^:?-+:?$', c) for c in rows[1]):
        rows = [rows[0]] + rows[2:]
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            add_run_with_inline(p, cell_text)
            if i == 0:
                for r in p.runs:
                    r.bold = True


def md_to_docx(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding='utf-8')
    doc = Document()
    # Style global : Calibri 11
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Lignes vides
        if not line.strip():
            i += 1
            continue

        # Séparateur ---
        if line.strip() == '---':
            p = doc.add_paragraph()
            p_run = p.add_run('─' * 40)
            p_run.font.color.rgb = MUTED
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Tables — détectées par début de ligne par '|'
        if line.strip().startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            add_table(doc, tbl_lines)
            doc.add_paragraph()
            continue

        # Headers
        if line.startswith('#'):
            m = re.match(r'^(#{1,6})\s+(.*)', line)
            if m:
                level = len(m.group(1))
                content = m.group(2).strip()
                # Retire ancres markdown éventuelles
                content = re.sub(r'\s*\{#[^}]+\}\s*$', '', content)
                h = doc.add_heading(level=min(level, 4))
                # Reset color & ensure visible (Word style headings are colored)
                for r in h.runs:
                    r.font.color.rgb = BRAND if level == 1 else None
                add_run_with_inline(h, content)
                i += 1
                continue

        # Listes - ou * ou 1.
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            content = list_match.group(3)
            style_name = 'List Number' if marker.endswith('.') else 'List Bullet'
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
            add_run_with_inline(p, content)
            if indent:
                p.paragraph_format.left_indent = Cm(0.5 * (indent // 2 + 1))
            i += 1
            continue

        # Blockquote / admonition  > [!TYPE] Titre
        if line.lstrip().startswith('>'):
            bq = []
            while i < len(lines) and lines[i].lstrip().startswith('>'):
                bq.append(re.sub(r'^\s*>\s?', '', lines[i].rstrip()))
                i += 1
            label = None
            if bq:
                am = re.match(r'^\[!(\w+)\]\s*(.*)$', bq[0])
                if am:
                    LABELS = {
                        'TIP': 'Bon réflexe', 'NOTE': 'À retenir', 'INFO': 'À retenir',
                        'IMPORTANT': 'Important', 'WARNING': 'Attention',
                        'DANGER': 'À éviter', 'SUCCESS': 'La promesse', 'CAUTION': 'À éviter',
                    }
                    label = am.group(2).strip() or LABELS.get(am.group(1).upper(), am.group(1).title())
                    bq[0] = ''
            if label:
                pl = doc.add_paragraph()
                rl = pl.add_run(label)
                rl.bold = True
                rl.font.color.rgb = BRAND
            for bl in bq:
                if not bl.strip():
                    continue
                lm = re.match(r'^(\s*)([-*+])\s+(.*)', bl)
                if lm:
                    try:
                        bp = doc.add_paragraph(style='List Bullet')
                    except KeyError:
                        bp = doc.add_paragraph()
                        bp.add_run('• ')
                    add_run_with_inline(bp, lm.group(3))
                else:
                    bp = doc.add_paragraph()
                    add_run_with_inline(bp, bl)
                bp.paragraph_format.left_indent = Cm(0.6)
            continue

        # Code block ```
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            r = p.add_run('\n'.join(code_lines))
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
            continue

        # Paragraphe normal — peut s'étendre sur plusieurs lignes consécutives
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].startswith('#')
            or lines[i].strip().startswith('|')
            or lines[i].lstrip().startswith(('- ', '* ', '+ '))
            or re.match(r'^\d+\.\s', lines[i].lstrip())
            or lines[i].strip() == '---'
            or lines[i].startswith('```')
            or lines[i].lstrip().startswith('>')
        ):
            para_lines.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        add_run_with_inline(p, ' '.join(para_lines))

    doc.save(out_path)
    print(f"OK : {out_path}")


if __name__ == '__main__':
    md_to_docx(SRC, DST)
    # Vincent 2026-05-22 — le mode d'emploi est aussi copié dans public/
    # pour servir au manager via /manuel. Les autres docs non.
    if BASENAME == "MODE_EMPLOI":
        PUBLIC_DST = ROOT / "public" / "manuel.docx"
        try:
            import shutil
            if DST.exists():
                shutil.copy2(DST, PUBLIC_DST)
                print(f"OK : copie servable {PUBLIC_DST}")
        except Exception as e:
            print(f"[WARN] copie public/manuel.docx echouee : {e}")
